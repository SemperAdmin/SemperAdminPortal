#!/usr/bin/env python3
"""Parse uploaded IGMC FAC docx files into JSON for inspectionSchema.

Optionally consume a canonical FA registry to override sponsor,
effectiveDate, category, faName, and title.
"""
import json, os, re, sys
from pathlib import Path
from datetime import datetime
from docx import Document

SUBSECTION_RE = re.compile(r"Subsection\s+(\d+)\s*[–—\-]+\s*(.*)", re.I)
ITEM_CODE_RE = re.compile(r"^\d{4}$")
REFERENCE_LEAD_RE = re.compile(r"^references?\s*:", re.I)
DATE_FORMATS = ["%d %B %Y", "%d %b %Y", "%B %d, %Y", "%b %d, %Y", "%d-%b-%y"]


def slugify(text):
    text = re.sub(r"\(.*?\)", "", text)
    text = re.sub(r"[^a-zA-Z0-9]+", "-", text).strip("-").lower()
    return text or "untitled"


def normalize_for_match(text):
    """Strip parens, punctuation, and whitespace for fuzzy comparison."""
    t = re.sub(r"\(.*?\)", "", text)
    t = re.sub(r"[^a-zA-Z0-9]+", " ", t).strip().lower()
    return t


def parse_date(value):
    s = re.sub(r"\s+", " ", value.strip())
    for fmt in DATE_FORMATS:
        try:
            return datetime.strptime(s, fmt).date().isoformat()
        except ValueError:
            continue
    return s


def first_match(prefix, text):
    pat = re.compile(r"^\s*" + re.escape(prefix) + r"\s*:\s*(.*)", re.I | re.M)
    m = pat.search(text)
    return m.group(1).strip() if m else None



DOC_START = re.compile(
    r"^(MCO|MARADMIN|MCBul|MCBUL|MCBUL\.|DODI|DODD|DODM|DODFMR|DoDFMR|DoD\s|NAVMC|SECNAV|SECNAVINST|SECNAV-M|MCTFSPRIUM|FPM|OPNAVINST|JTR|JFTR|UCMJ|ALMAR|MARDIV|JAGMAN|CFR|USC|MCRP|MCDP|MCWP|TM|NAVPERS|NAVADMIN|FMR|JAGINST|JAGMAN|MARCORSEPMAN|MIL-STD|MIL-HDBK|JTR|JFTR|MARFORRES|MCRP|SOP|MCASOP)",
    re.I,
)


def is_doc_start(token):
    return bool(DOC_START.match(token.strip()))


def carry_forward(parts):
    """Merge continuation fragments into their preceding doc citation."""
    out = []
    for p in parts:
        p = p.strip().rstrip(".")
        if not p:
            continue
        if out and not is_doc_start(p):
            out[-1] = out[-1] + "; " + p
        else:
            out.append(p)
    return out

def split_references(line):
    body = re.sub(r"^references?\s*:\s*", "", line, flags=re.I).strip()
    parts = re.split(r"\s*;\s*", body)
    return carry_forward(parts)


def walk_tables(tables, out):
    for table in tables:
        seen_rows = set()
        for row in table.rows:
            tr_id = id(row._tr)
            if tr_id in seen_rows:
                continue
            seen_rows.add(tr_id)
            cells = list(row.cells)
            texts = [c.text.strip() for c in cells]
            out.append((texts, cells))
            seen_n = set()
            for cell in cells:
                for nt in cell.tables:
                    nid = id(nt._tbl)
                    if nid in seen_n:
                        continue
                    seen_n.add(nid)
                    walk_tables([nt], out)


def collect_rows(doc):
    out = []
    walk_tables(doc.tables, out)
    return out


def extract_metadata(rows):
    meta = {"title_raw": "", "applicability": "", "sponsor": "", "sme": "",
            "revised": "", "total_questions": None}
    if not rows:
        return meta
    for t in rows[0][0]:
        if t:
            meta["title_raw"] = t
            break
    seen = set()
    flat = []
    for texts, _ in rows:
        for t in texts:
            if t and t not in seen:
                seen.add(t)
                flat.append(t)
                if SUBSECTION_RE.match(t) or ITEM_CODE_RE.match(t):
                    break
    for t in flat:
        if re.match(r"^\s*this checklist (applies|is applicable)", t, re.I):
            cleaned = re.sub(r"^\s*this checklist (applies\s+to|is\s+applicable\s+to)\s*",
                             "", t, flags=re.I).strip()
            cleaned = re.sub(r"\s+", " ", cleaned).rstrip(".")
            if cleaned:
                cleaned = cleaned[0].upper() + cleaned[1:]
            meta["applicability"] = cleaned
            break
    big = "\n".join(flat)
    for key, label in [("sponsor", "Functional Area Sponsor"),
                       ("sme", "Subject Matter Expert"),
                       ("revised", "Revised")]:
        v = first_match(label, big)
        if v:
            meta[key] = v
    tq = first_match("Total Questions", big)
    if tq:
        m = re.search(r"\d+", tq)
        if m:
            meta["total_questions"] = int(m.group(0))
    return meta


def parse_item_cell(code, cell):
    if cell is None:
        return code, "", [], None
    paras = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    refs = []
    qlines = []
    evidence = None
    for p in paras:
        if REFERENCE_LEAD_RE.match(p):
            refs.extend(split_references(p))
            continue
        if re.match(r"^\s*evidence\s*:", p, re.I):
            evidence = re.sub(r"^\s*evidence\s*:\s*", "", p, flags=re.I).strip()
            continue
        qlines.append(p)
    q = re.sub(r"\s+", " ", " ".join(qlines)).strip()
    inline = re.search(r"references?\s*:\s*(.+)$", q, re.I)
    if inline and not refs:
        refs = split_references(inline.group(0))
        q = q[:inline.start()].strip()
    return code, q, refs, evidence


def extract_subsections(rows):
    subs = []
    cur = None
    for texts, cells in rows:
        if not texts:
            continue
        c0 = texts[0]
        sm = SUBSECTION_RE.match(c0)
        if sm:
            if cur:
                subs.append(cur)
            best = sm.group(2).strip()
            for t in texts:
                m = SUBSECTION_RE.match(t)
                if m and len(m.group(2)) > len(best):
                    best = m.group(2).strip()
            cur = {"id": sm.group(1).zfill(2) + ".0", "title": best,
                   "description": "", "items": []}
            continue
        if ITEM_CODE_RE.match(c0):
            if cur is None:
                cur = {"id": "01.0", "title": "Items", "description": "", "items": []}
            qcell = None
            for idx in range(1, len(cells)):
                if texts[idx] and texts[idx] != c0:
                    qcell = cells[idx]
                    break
            code, q, refs, ev = parse_item_cell(c0, qcell)
            if not q:
                continue
            item = {"code": code, "question": q, "references": refs}
            if ev:
                item["evidenceHint"] = ev
            cur["items"].append(item)
    if cur:
        subs.append(cur)
    return subs


def split_by_code_prefix_if_needed(subs):
    if len(subs) != 1:
        return subs
    sub = subs[0]
    if not sub["items"]:
        return subs
    prefixes = {it["code"][:2] for it in sub["items"]}
    if len(prefixes) == 1:
        return subs
    by = {}
    order = []
    for it in sub["items"]:
        pre = it["code"][:2]
        if pre not in by:
            by[pre] = []
            order.append(pre)
        by[pre].append(it)
    out = []
    for i, pre in enumerate(order):
        out.append({"id": pre + ".0",
                    "title": sub["title"] if i == 0 else "Section " + str(int(pre)),
                    "description": sub["description"] if i == 0 else "",
                    "items": by[pre]})
    return out


def derive_program_number(title_raw, filename):
    m = re.search(r"\b(\d{3,4}(?:\.\d+)?)\b", filename)
    if m:
        return m.group(1)
    m = re.search(r"\((\d{3,4}(?:\.\d+)?)\)\s*$", title_raw)
    if m:
        return m.group(1)
    m = re.search(r"\b(\d{3,4}(?:\.\d+)?)\b", title_raw)
    return m.group(1) if m else filename


def derive_title(title_raw, filename):
    if not title_raw or re.match(r"^\s*Functional Area Sponsor", title_raw, re.I):
        cleaned = re.sub(r"\.docx?$", "", filename, flags=re.I)
        cleaned = re.sub(r"\s*\(\d+\)\s*", " ", cleaned)
        cleaned = re.sub(r"\b\d{3,4}(?:\.\d+)?\b", "", cleaned)
        cleaned = re.sub(r"revised.*$", "", cleaned, flags=re.I)
        cleaned = re.sub(r"\s+", " ", cleaned).strip(" -")
        return cleaned.title() or "Untitled"
    cleaned = re.sub(r"\s*\(\d{3,4}(?:\.\d+)?(?:,\s*Chapter\s+\d+)?\)\s*$", "", title_raw)
    cleaned = re.sub(r"\s+\d{3,4}(?:\.\d+)?\s*$", "", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip().title()


def derive_summary(applicability, total_q):
    base = applicability or "Inspector General Functional Area Checklist."
    base = re.sub(r"\s+", " ", base).strip()
    if total_q:
        base = base.rstrip(".") + ". " + str(total_q) + " items across the program."
    if len(base) > 280:
        base = base[:275].rstrip() + "..."
    return base


def load_registry(path):
    """Load the FA registry and key by faNumber. Each entry has a list
    of canonical FAs that share the number (e.g., 5800.16)."""
    if not path or not Path(path).exists():
        return {}
    data = json.loads(Path(path).read_text())
    bucket = {}
    for entry in data.get("entries", []):
        bucket.setdefault(entry["faNumber"], []).append(entry)
    return bucket


def lookup_registry(registry, program_number, parsed_title):
    """Return the best registry entry for a parsed record, or None."""
    candidates = registry.get(program_number, [])
    if not candidates:
        # Try the FA-level number when a parsed file used the MCO-revision number.
        fa_only = re.sub(r"\.\d+$", "", program_number)
        candidates = registry.get(fa_only, [])
    if len(candidates) == 1:
        return candidates[0]
    if not candidates:
        return None
    # Multiple candidates share the FA number. Match by title.
    pt = normalize_for_match(parsed_title)
    best = None
    best_score = -1
    for c in candidates:
        ft = normalize_for_match(c["faName"])
        # Score: number of overlapping word tokens.
        pt_tokens = set(pt.split())
        ft_tokens = set(ft.split())
        overlap = len(pt_tokens & ft_tokens)
        if overlap > best_score:
            best_score = overlap
            best = c
    return best


def build_record(path, registry):
    doc = Document(str(path))
    rows = collect_rows(doc)
    meta = extract_metadata(rows)
    subs = extract_subsections(rows)
    program_number = derive_program_number(meta["title_raw"], path.name)
    title = derive_title(meta["title_raw"], path.name) or path.stem
    revised_iso = parse_date(meta["revised"]) if meta["revised"] else "2025-01-01"
    sponsor = meta["sponsor"] or "Unspecified"
    category = None
    fa_name = None
    effective_iso = revised_iso

    reg = lookup_registry(registry, program_number, title)
    if reg:
        sponsor = reg["sponsor"]
        category = reg["category"]
        fa_name = reg["faName"]
        effective_iso = reg["effectiveDate"]
        # Use the canonical FA name as the title for consistency.
        title = reg["faName"]
        program_number = reg["faNumber"]

    slug = slugify(title)
    subs = split_by_code_prefix_if_needed(subs)
    summary = derive_summary(meta["applicability"], meta["total_questions"])

    record = {
        "track": "igmc",
        "programNumber": program_number,
        "slug": slug,
        "title": title,
        "summary": summary,
        "roles": ["admin", "commander"],
        "sponsor": sponsor,
        "applicabilityLevel": meta["applicability"] or None,
        "effectiveDate": effective_iso,
        "lastVerified": revised_iso,
        "source": {
            "title": "IGMC Functional Area Checklist - " + title,
            "publisher": "Inspector General of the Marine Corps",
        },
        "subsections": subs,
    }
    if category:
        record["category"] = category
    if fa_name:
        record["faName"] = fa_name
    if not record["applicabilityLevel"]:
        del record["applicabilityLevel"]
    if sum(len(s["items"]) for s in subs) == 0:
        raise SystemExit("NO ITEMS FOUND for " + path.name)
    return record


def main():
    args = sys.argv[1:]
    registry_path = None
    if "--registry" in args:
        i = args.index("--registry")
        registry_path = args[i + 1]
        del args[i:i + 2]
    if len(args) != 2:
        print("Usage: parse_igmc.py [--registry path] <uploads-dir> <output-dir>",
              file=sys.stderr)
        return 2
    src = Path(args[0])
    dst = Path(args[1])
    dst.mkdir(parents=True, exist_ok=True)
    registry = load_registry(registry_path)
    files = sorted(src.glob("*.DOCX")) + sorted(src.glob("*.docx"))
    written = []
    skipped = []
    unmatched = []
    for path in files:
        try:
            record = build_record(path, registry)
        except Exception as exc:
            skipped.append((path.name, str(exc)))
            continue
        if registry and "category" not in record:
            unmatched.append((path.name, record["programNumber"], record["title"]))
        out_path = dst / (record["programNumber"] + "-" + record["slug"] + ".json")
        out_path.write_text(json.dumps(record, indent=2, ensure_ascii=False))
        items_n = sum(len(s["items"]) for s in record["subsections"])
        written.append((path.name, out_path.name, len(record["subsections"]), items_n))
    print("Parsed", len(written), "files,", len(skipped), "skipped")
    for w in written:
        print("  OK  " + w[0] + " -> " + w[1] + " (" + str(w[2]) + " subs, " + str(w[3]) + " items)")
    for s in skipped:
        print("  ERR " + s[0] + ": " + s[1])
    if unmatched:
        print("\nNo registry match for:")
        for u in unmatched:
            print("  " + u[1] + " (" + u[2] + ")  from " + u[0])
    return 0 if not skipped else 1


if __name__ == "__main__":
    sys.exit(main())
